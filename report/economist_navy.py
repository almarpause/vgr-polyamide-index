# -*- coding: utf-8 -*-
"""
economist_navy.py — Economist-Word-Style (Navy edition) build library.
Reusable helpers that impose The Economist's print grammar on a .docx,
re-skinned so Navy replaces Econ Red as the single structural accent.
python-docx based. See economist-word-style-SPEC.md for the full spec.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Emu, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- palette ----------
ACCENT_DK = "C0110A"   # deep Econ red (large fills)
NAVY900 = "0D1F3C"     # secondary/data navy
NAVY700 = "12345B"     # secondary/data navy
ACCENT    = "E3120B"   # Econ Red (structural accent, marks on white)
STEEL   = "1F3F6B"
SKY     = "4A7BAE"
PANEL   = "E8EDF2"
COOL    = "F4F6F9"
RULEGREY= "C5CDD6"
MIDGREY = "5A6A7A"
INK     = "1A1A1A"
SIGNALRED = "E3120B"

SERIF = "Georgia"
SANS  = "Arial"

def _rgb(h): return RGBColor.from_string(h)

# ---------- low-level XML helpers ----------
def _set(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    el.append(e)
    return e

def para_border(p, edges=("bottom",), color=ACCENT, size=12, space=6, val="single"):
    """size in eighths of a point (12 = 1.5pt). edges: top/bottom/left/right."""
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    for edge in edges:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), val); e.set(qn('w:sz'), str(size))
        e.set(qn('w:space'), str(space)); e.set(qn('w:color'), color)
        pbdr.append(e)

def para_shading(p, fill=PANEL):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for k, v in (("top",top),("bottom",bottom),("start",left),("end",right),("left",left),("right",right)):
        e = OxmlElement(f'w:{k}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)

def set_cell_vertical(cell, align="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), align); tcPr.append(va)

def no_table_borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'none'); e.set(qn('w:sz'),'0')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'auto'); borders.append(e)
    tblPr.append(borders)

def ledger_borders(table, top_color=ACCENT_DK, top_sz=14, bottom_color=ACCENT_DK,
                   bottom_sz=14, hair_color=RULEGREY, hair_sz=4, insideH=True):
    tbl = table._tbl; tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    spec = {
        'top':    (top_color, top_sz, 'single'),
        'bottom': (bottom_color, bottom_sz, 'single'),
        'left':   ('auto', 0, 'none'),
        'right':  ('auto', 0, 'none'),
        'insideH':(hair_color, hair_sz, 'single' if insideH else 'none'),
        'insideV':('auto', 0, 'none'),
    }
    for edge,(c,s,v) in spec.items():
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),v); e.set(qn('w:sz'),str(s))
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),c); borders.append(e)
    tblPr.append(borders)

def run_tracking(run, twips=8):
    """character spacing (letter-spacing), value in twips (20ths of a pt)."""
    rPr = run._r.get_or_add_rPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(twips)); rPr.append(sp)

def keep_with_next(p):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement('w:keepNext'); pPr.append(e)
    e2 = OxmlElement('w:keepLines'); pPr.append(e2)

def enable_hyphenation(doc):
    settings = doc.settings.element
    for tag in ('w:autoHyphenation','w:doNotHyphenateCaps'):
        ex = settings.find(qn(tag))
        if ex is None:
            e = OxmlElement(tag); e.set(qn('w:val'), 'true' if tag=='w:autoHyphenation' else 'false')
            settings.append(e)


def _row_header(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)

def _row_cantsplit(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement('w:cantSplit'); e.set(qn('w:val'),'true'); trPr.append(e)

# ---------- run/paragraph builders ----------
def add_run(p, text, font=SERIF, size=10.5, color=INK, bold=False, italic=False,
            caps=False, track=None):
    r = p.add_run(text.upper() if caps else text)
    r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = _rgb(color)
    r.font.bold = bold; r.font.italic = italic
    # ensure east-asian/complex fallback uses same font
    rPr = r._r.get_or_add_rPr(); rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    for a in ('w:ascii','w:hAnsi','w:cs'): rFonts.set(qn(a), font)
    if track: run_tracking(r, track)
    return r

def _new_para(doc, space_before=0, space_after=6, align=None, line=1.28, left_indent=None,
              right_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    if align is not None: pf.alignment = align
    if line: pf.line_spacing = line
    if left_indent is not None: pf.left_indent = Inches(left_indent)
    if right_indent is not None: pf.right_indent = Inches(right_indent)
    return p


def _add_cell_text(p, val, font, size, base_color, bold=False):
    """Render cell text, coloring [[r]]...[[/r]] spans in SIGNALRED."""
    import re as _re
    val = str(val)
    parts = _re.split(r'(\[\[r\]\].*?\[\[/r\]\])', val)
    for part in parts:
        if not part:
            continue
        if part.startswith('[[r]]') and part.endswith('[[/r]]'):
            add_run(p, part[5:-6], font=font, size=size, color=SIGNALRED, bold=True)
        else:
            add_run(p, part, font=font, size=size, color=base_color, bold=bold)

# ---------- components ----------
def kicker(doc, text):
    p = _new_para(doc, space_before=14, space_after=2, line=1.0)
    para_border(p, edges=("top",), color=ACCENT, size=12, space=6)
    keep_with_next(p)
    add_run(p, text, font=SANS, size=9, color=ACCENT, bold=True, caps=True, track=10)
    return p

def h1(doc, text, deck=None):
    p = _new_para(doc, space_before=6, space_after=2, line=1.02)
    keep_with_next(p)
    add_run(p, text, font=SANS, size=22, color=INK, bold=True)
    if deck:
        d = _new_para(doc, space_before=2, space_after=10, line=1.15)
        add_run(d, deck, font=SERIF, size=12, color=MIDGREY, italic=True)
    return p

def h2(doc, text):
    p = _new_para(doc, space_before=2, space_after=4, line=1.05)
    keep_with_next(p)
    add_run(p, text, font=SANS, size=13, color=INK, bold=True)
    return p

def h3(doc, text):
    p = _new_para(doc, space_before=8, space_after=2, line=1.05)
    keep_with_next(p)
    add_run(p, text, font=SANS, size=10.5, color=INK, bold=True)
    return p

def body(doc, text, drop=False, space_after=6):
    p = _new_para(doc, space_after=space_after, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.28)
    if drop and text:
        # true drop cap: first letter in a framed paragraph
        _drop_cap_letter(doc, p, text[0])
        add_run(p, text[1:], font=SERIF, size=10.5, color=INK)
    else:
        add_run(p, text, font=SERIF, size=10.5, color=INK)
    return p

def _drop_cap_letter(doc, host_para, letter):
    """Insert a proper dropCap frame paragraph before host_para's text."""
    dp = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
    frame = OxmlElement('w:framePr')
    frame.set(qn('w:dropCap'),'drop'); frame.set(qn('w:lines'),'3')
    frame.set(qn('w:wrap'),'around'); frame.set(qn('w:vAnchor'),'text')
    frame.set(qn('w:hAnchor'),'text'); frame.set(qn('w:h'),'1080'); frame.set(qn('w:w'),'720')
    pPr.append(frame)
    spac = OxmlElement('w:spacing'); spac.set(qn('w:line'),'240'); spac.set(qn('w:lineRule'),'exact')
    pPr.append(spac)
    dp.append(pPr)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), SERIF)
    rPr.append(rf)
    b = OxmlElement('w:sz'); b.set(qn('w:val'), str(int(52*2))); rPr.append(b)
    col = OxmlElement('w:color'); col.set(qn('w:val'), ACCENT); rPr.append(col)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = letter; r.append(t)
    dp.append(r)
    host_para._p.addprevious(dp)

def bullets(doc, items, numbered=False):
    for i, it in enumerate(items):
        p = _new_para(doc, space_after=3, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.22,
                      left_indent=0.3)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        marker = f"{i+1}  " if numbered else "—  "
        mr = add_run(p, marker, font=SANS, size=10.5, color=ACCENT, bold=True)
        add_run(p, it, font=SERIF, size=10.5, color=INK)

def pull_quote(doc, text, attrib=None):
    p = _new_para(doc, space_before=10, space_after=10, line=1.2, left_indent=0.3, right_indent=0.5)
    para_border(p, edges=("left",), color=ACCENT, size=18, space=10)
    add_run(p, text, font=SERIF, size=15, color=INK, italic=False)
    if attrib:
        a = _new_para(doc, space_before=0, space_after=10, left_indent=0.3)
        add_run(a, "— "+attrib, font=SANS, size=8.5, color=MIDGREY, caps=True, track=6)

def source_line(doc, text):
    p = _new_para(doc, space_before=2, space_after=10, line=1.0)
    add_run(p, text, font=SANS, size=8, color=MIDGREY)

def stat_callout(doc, number, label):
    p = _new_para(doc, space_before=6, space_after=0, line=1.0)
    add_run(p, number, font=SANS, size=34, color=ACCENT_DK, bold=True)
    l = _new_para(doc, space_before=0, space_after=8, line=1.0)
    add_run(l, label, font=SANS, size=8.5, color=MIDGREY, caps=True, track=8)

def callout(doc, label, lines):
    # panel: shaded paragraph(s) with a navy top bar
    top = _new_para(doc, space_before=10, space_after=0, line=0.3)
    para_border(top, edges=("top",), color=ACCENT_DK, size=24, space=0)
    para_shading(top, PANEL)
    p = _new_para(doc, space_before=0, space_after=0, line=1.2, left_indent=0.12, right_indent=0.12)
    para_shading(p, PANEL)
    add_run(p, label, font=SANS, size=8.5, color=ACCENT, bold=True, caps=True, track=8)
    for i, ln in enumerate(lines):
        q = _new_para(doc, space_before=1, space_after=(8 if i==len(lines)-1 else 1),
                      line=1.22, left_indent=0.12, right_indent=0.12)
        para_shading(q, PANEL)
        add_run(q, ln, font=SERIF, size=10, color=INK)

def ledger_table(doc, headers, rows, widths, source=None, zebra=True,
                 right_cols=(), red_negatives=False):
    total = sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    # column widths
    tbl = table._tbl
    grid = tbl.tblGrid
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(w))
    # table width
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(total)); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    ledger_borders(table)
    # header
    hdr = table.rows[0].cells
    for j, htext in enumerate(headers):
        set_cell_margins(hdr[j]); set_cell_vertical(hdr[j], "bottom")
        hdr[j].width = Twips(widths[j])
        p = hdr[j].paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing=1.0
        if j in right_cols: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, htext, font=SANS, size=8.5, color=ACCENT_DK, bold=True, caps=True, track=4)
    _row_header(table.rows[0]); _row_cantsplit(table.rows[0])
    # body rows
    for ri, row in enumerate(rows):
        _brow = table.add_row(); cells = _brow.cells; _row_cantsplit(_brow)
        for j, val in enumerate(row):
            set_cell_margins(cells[j]); set_cell_vertical(cells[j], "center")
            cells[j].width = Twips(widths[j])
            if zebra and ri % 2 == 1:
                cell_shading(cells[j], COOL)
            p = cells[j].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.line_spacing=1.08
            is_num = j in right_cols
            if is_num: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            col = INK
            if red_negatives and is_num and str(val).strip().startswith(('-','−','▼')):
                col = SIGNALRED
            _add_cell_text(p, val, font=(SANS if is_num else SERIF), size=9.5,
                           base_color=col, bold=(j==0 and not is_num))
    if source:
        source_line(doc, source)
    return table


def figure(doc, png_path, width_in=6.3, space_before=8, space_after=6):
    p = _new_para(doc, space_before=space_before, space_after=space_after, line=1.0)
    r = p.add_run()
    r.add_picture(png_path, width=Inches(width_in))
    return p

def contents_row(doc, num, label, page, measure_in=6.3):
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.enum.text import WD_TAB_LEADER
    p = _new_para(doc, space_after=3, line=1.15)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(measure_in), WD_TAB_ALIGNMENT.RIGHT,
                                              WD_TAB_LEADER.DOTS)
    add_run(p, num+"   ", font=SANS, size=9, color=ACCENT, bold=True)
    add_run(p, label, font=SERIF, size=10.5, color=INK)
    add_run(p, "\t", font=SERIF, size=10.5, color=INK)
    add_run(p, str(page), font=SANS, size=9, color=MIDGREY)
    return p

def colophon(doc, lines, space_before=16):
    p = _new_para(doc, space_before=space_before, space_after=4, line=1.0)
    para_border(p, edges=("top",), color=ACCENT, size=10, space=8)
    for i, (txt, kind) in enumerate(lines):
        q = _new_para(doc, space_before=(0 if i==0 else 1), space_after=1, line=1.2)
        if kind == "label":
            add_run(q, txt, font=SANS, size=8.5, color=ACCENT, bold=True, caps=True, track=8)
        else:
            add_run(q, txt, font=SERIF, size=9, color=MIDGREY)
    return p

# ---------- cover / footer ----------
def cover_masthead(doc, label):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    w = 2160
    t._tbl.tblGrid.findall(qn('w:gridCol'))[0].set(qn('w:w'), str(w))
    tblPr = t._tbl.tblPr
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(w)); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    no_table_borders(t)
    c = t.rows[0].cells[0]; c.width = Twips(w)
    cell_shading(c, ACCENT_DK); set_cell_margins(c, 70, 70, 130, 130)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing=1.0
    add_run(p, label, font=SANS, size=13, color="FFFFFF", bold=True, caps=True, track=10)
    return t

def stripe_band(doc, hexes=None):
    if hexes is None:
        hexes = [ACCENT_DK, "10294B", STEEL, "2A5488", SKY, "6E9BC6", "9DBBDA", PANEL]
    n = len(hexes)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    total = 9060; cw = total // n
    for gc in t._tbl.tblGrid.findall(qn('w:gridCol')): gc.set(qn('w:w'), str(cw))
    tblPr = t._tbl.tblPr
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(cw*n)); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    no_table_borders(t)
    for j, hx in enumerate(hexes):
        cel = t.rows[0].cells[j]; cel.width = Twips(cw)
        cell_shading(cel, hx); set_cell_margins(cel, 0,0,0,0)
        p = cel.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
        add_run(p, " ", font=SANS, size=9, color=hx)
    return t

def add_page_number_field(paragraph):
    r = paragraph.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    rr = OxmlElement('w:r'); rp = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii','w:hAnsi'): rf.set(qn(a), SERIF)
    rp.append(rf); sz=OxmlElement('w:sz'); sz.set(qn('w:val'),'16'); rp.append(sz)
    col=OxmlElement('w:color'); col.set(qn('w:val'),MIDGREY); rp.append(col)
    rr.append(rp); tt=OxmlElement('w:t'); tt.text='1'; rr.append(tt); fld.append(rr)
    paragraph._p.append(fld)

def setup_footer(section, short_title):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.line_spacing = 1.0
    para_border(p, edges=("top",), color=RULEGREY, size=8, space=6)
    # left title + right page number via tab stop
    from docx.enum.text import WD_TAB_ALIGNMENT
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    add_run(p, short_title, font=SANS, size=8, color=MIDGREY, caps=True, track=6)
    add_run(p, "\t", font=SANS, size=8, color=MIDGREY)
    add_page_number_field(p)

def new_document(short_title="", left_margin=1.1, right_margin=1.1, top=1.0, bottom=0.9):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = Inches(left_margin); sec.right_margin = Inches(right_margin)
    sec.top_margin = Inches(top); sec.bottom_margin = Inches(bottom)
    # default style
    normal = doc.styles['Normal']
    normal.font.name = SERIF; normal.font.size = Pt(10.5); normal.font.color.rgb = _rgb(INK)
    enable_hyphenation(doc)
    if short_title:
        setup_footer(sec, short_title)
    return doc

def content_width_twips(doc):
    sec = doc.sections[0]
    return int((sec.page_width - sec.left_margin - sec.right_margin) / 635)  # EMU->twips

def page_break(doc):
    doc.add_page_break()

def spacer(doc, pts=6):
    p = _new_para(doc, space_before=0, space_after=0, line=1.0)
    add_run(p, "", font=SERIF, size=pts)
